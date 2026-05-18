from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class ParsedPage:
    text: str
    page_number: int | None = None


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    pages: list[ParsedPage]


def extract_text_from_file(path: Path, mime_type: str) -> str:
    return parse_file(path, mime_type).text


def parse_file(path: Path, mime_type: str) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf" or mime_type == "application/pdf":
        return _parse_pdf(path)
    if suffix in (".docx",) or mime_type.endswith("wordprocessingml.document"):
        return _parse_docx(path)
    if suffix in (".txt", ".md"):
        text = _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
        return ParsedDocument(text=text, pages=[ParsedPage(text=text)])
    raise ValueError(f"Unsupported file type: {suffix}")


def _parse_pdf(path: Path) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [
        ParsedPage(text=_normalize_text(page.extract_text() or ""), page_number=index)
        for index, page in enumerate(reader.pages, start=1)
    ]
    text = "\n\n".join(page.text for page in pages if page.text)
    return ParsedDocument(text=text, pages=[page for page in pages if page.text])


def _parse_docx(path: Path) -> ParsedDocument:
    from docx import Document

    doc = Document(str(path))
    text = "\n\n".join(_normalize_text(p.text) for p in doc.paragraphs if p.text.strip())
    return ParsedDocument(text=text, pages=[ParsedPage(text=text)])


def _normalize_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())
